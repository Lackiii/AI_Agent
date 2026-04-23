"""
Merge tables and images from API design docx into thesis draft docx.
Preserves block order (paragraphs + tables). Uses deep copy of OOXML elements.

Source: docs/API接口与模块交互设计.docx (Word lock files ~$.docx are ignored).
Insertion: before heading \"3.6 本章小结\" in chapter 3 (after 3.5 数据与交互说明).
"""
from __future__ import annotations

import os
import re
import shutil
import sys
from copy import deepcopy

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


def _find_docs_dir() -> str:
    root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    return os.path.join(root, "docs")


def _resolve_paths(docs_dir: str) -> tuple[str, str]:
    thesis = None
    api_src = None
    for name in os.listdir(docs_dir):
        low = name.lower()
        if low.startswith("~$"):
            continue
        if low.endswith(".docx"):
            if low.startswith("api"):
                api_src = os.path.join(docs_dir, name)
            if "2022433020102" in name and "backup" not in low:
                thesis = os.path.join(docs_dir, name)
    if not api_src:
        for name in os.listdir(docs_dir):
            if name.lower().startswith("~$"):
                continue
            if name.lower().endswith(".docx") and name.lower().startswith("api"):
                api_src = os.path.join(docs_dir, name)
    if not thesis:
        for name in os.listdir(docs_dir):
            low = name.lower()
            if low.startswith("~$"):
                continue
            if not low.endswith(".docx"):
                continue
            if low.startswith("api"):
                continue
            if "merged" in low or "backup" in low:
                continue
            thesis = os.path.join(docs_dir, name)
            break
    if not thesis or not api_src:
        raise FileNotFoundError(f"thesis={thesis!r} api_src={api_src!r} in {docs_dir}")
    return thesis, api_src


def iter_block_items(doc: DocumentType):
    parent_elm = doc.element.body
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def paragraph_has_image(p: Paragraph) -> bool:
    for run in p.runs:
        el = run._element
        if el.findall(".//" + qn("w:drawing")) or el.findall(".//" + qn("w:pict")):
            return True
    return False


def is_effectively_empty_paragraph(p: Paragraph) -> bool:
    t = (p.text or "").strip()
    return len(t) == 0 and not paragraph_has_image(p)


def collect_blocks(doc: DocumentType):
    """Return list of (kind, obj) in document order: kind in ('p','tbl')."""
    out = []
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            out.append(("p", block))
        else:
            out.append(("tbl", block))
    return out


def filter_blocks(blocks):
    """Keep tables and paragraphs that have text or images; drop empty paragraphs."""
    kept = []
    for kind, obj in blocks:
        if kind == "tbl":
            kept.append((kind, obj))
            continue
        if is_effectively_empty_paragraph(obj):
            continue
        kept.append((kind, obj))
    return kept


def _style_name(p: Paragraph) -> str:
    try:
        return p.style.name if p.style else ""
    except (KeyError, AttributeError):
        return ""


def find_insert_before_3_6_heading(doc: DocumentType) -> Paragraph | None:
    """Insert API blocks immediately before '3.6 本章小结'."""
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        st = _style_name(p)
        if not st.startswith("Heading"):
            continue
        if t.startswith("3.6") and "小结" in t:
            return p
    return None


def find_insert_after_module_summary(doc: DocumentType) -> Paragraph | None:
    """Fallback: after the paragraph that summarizes cross-module interaction."""
    pat = re.compile(r"模块.*交互关系.*总结")
    last_match: Paragraph | None = None
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if pat.search(text):
            last_match = p
    return last_match


def insert_before_paragraph(target: DocumentType, p: Paragraph, elements: list):
    parent = target.element.body
    children = list(parent.iterchildren())
    idx = children.index(p._element)
    for i, el in enumerate(elements):
        parent.insert(idx + i, el)


def build_heading_elements(doc: DocumentType, body) -> list:
    """Section title + bridging paragraph; match thesis Heading 3 / Normal."""
    elements = []
    h = doc.add_heading("3.5.1 接口与模块交互设计", level=3)
    elements.append(deepcopy(h._element))
    body.remove(h._element)

    bridge = doc.add_paragraph(
        "为与上文总体架构及数据流描述相呼应，本节以接口列表、模块交互图与数据表形式，"
        "给出主要 REST/WebSocket/IPC 接口及模块间调用关系（内容与图表来源于《API 接口与模块交互设计》）。"
    )
    elements.append(deepcopy(bridge._element))
    body.remove(bridge._element)
    return elements


def main() -> int:
    docs_dir = _find_docs_dir()
    thesis_path, api_path = _resolve_paths(docs_dir)

    api_doc = Document(api_path)
    blocks = filter_blocks(collect_blocks(api_doc))

    thesis = Document(thesis_path)
    body = thesis.element.body

    anchor_before = find_insert_before_3_6_heading(thesis)
    anchor_after = None if anchor_before else find_insert_after_module_summary(thesis)

    to_insert: list = []
    to_insert.extend(build_heading_elements(thesis, body))

    for _kind, obj in blocks:
        to_insert.append(deepcopy(obj._element))

    if anchor_before is not None:
        insert_before_paragraph(thesis, anchor_before, to_insert)
    elif anchor_after is not None:
        children = list(body.iterchildren())
        idx = children.index(anchor_after._element) + 1
        for i, el in enumerate(to_insert):
            body.insert(idx + i, el)
    else:
        for el in to_insert:
            body.append(el)

    backup = thesis_path + ".backup_before_api_merge.docx"
    if not os.path.isfile(backup):
        shutil.copy2(thesis_path, backup)

    thesis.save(thesis_path)
    print("Updated thesis:", thesis_path)
    print("Backup (first run only):", backup)
    print("Source API doc:", api_path)
    return 0


if __name__ == "__main__":
    sys.exit(main())
