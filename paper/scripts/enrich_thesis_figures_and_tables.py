# -*- coding: utf-8 -*-
"""
Enrich thesis docx: move 表3-1 above table, unify captions, add figure/table explanations
and section transitions (style aligned with common graduation-thesis practice; reference:
paper/docs/4.参考论文1（软件开发类）.doc).

Run: python paper/scripts/enrich_thesis_figures_and_tables.py
"""
from __future__ import annotations

import os
import re
import sys
from copy import deepcopy

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

FIG_CAP_RE = re.compile(r"^图\s*(\d+)\s*-\s*(\d+)\s*[　 ]*(.+)$")


def _paper_root() -> str:
    return os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))


def _resolve_thesis_path(docs_dir: str) -> str:
    for name in os.listdir(docs_dir):
        if not name.endswith(".docx") or name.startswith("~$"):
            continue
        low = name.lower()
        if "2022433020102" not in name:
            continue
        if "backup" in low or "before_structure" in low:
            continue
        return os.path.join(docs_dir, name)
    raise FileNotFoundError(f"No main thesis docx in {docs_dir}")


def _para_text(el) -> str:
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    return "".join(t.text or "" for t in el.findall(".//w:t", namespaces=ns))


def _set_run_body_font(run):
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _body_paragraph(doc: Document, text: str, *, indent: bool = True) -> object:
    p = doc.add_paragraph(text)
    fmt = p.paragraph_format
    if indent:
        fmt.first_line_indent = Cm(0.74)
    fmt.line_spacing = 1.5
    for r in p.runs:
        _set_run_body_font(r)
    elm = deepcopy(p._element)
    body = doc.element.body
    body.remove(p._element)
    return elm


def _insert_before(body, anchor, elements: list):
    children = list(body.iterchildren())
    idx = children.index(anchor)
    for i, el in enumerate(elements):
        body.insert(idx + i, el)


def _insert_after(body, anchor, elements: list):
    children = list(body.iterchildren())
    idx = children.index(anchor) + 1
    for i, el in enumerate(elements):
        body.insert(idx + i, el)


def _normalize_figure_caption_text(raw: str) -> str | None:
    m = FIG_CAP_RE.match(raw.strip())
    if not m:
        return None
    ch, num, title = m.group(1), m.group(2), m.group(3).strip()
    return f"图{ch}-{num} {title}"


def EXPLANATIONS() -> dict[str, str]:
    return {
        "3-1": (
            "如图3-1所示，应用启动阶段以 Electron 主进程为入口完成环境加载、窗口创建与关键服务注册；"
            "该主链路强调“先完成安全边界与通信能力，再进入业务编排”的顺序，为后续渲染进程接入与后端服务发现奠定基础。"
        ),
        "3-2": (
            "如图3-2所示，渲染进程通过预加载脚本暴露的受控 API 与主进程建立 IPC 调用关系；"
            "该机制在不向页面开放 Node 高权限对象的前提下完成对话、记忆、提醒、截图等能力转发，体现了桌面应用最小权限原则。"
        ),
        "3-3": (
            "如图3-3所示，系统采用表示层—编排层—服务层的分层组织方式，各层通过清晰边界耦合；"
            "该结构有助于在需求变更时局部替换实现（例如替换模型服务或扩展后端能力）而不牵动全栈。"
        ),
        "3-4": (
            "如图3-4所示，核心对话子系统覆盖发送消息、管理记忆与人设、资料夹工具调用等典型参与者交互；"
            "用例边界明确了“用户—前端—主进程—外部模型/工具”的职责划分，为第4章实现提供可验证的需求基线。"
        ),
        "3-5": (
            "如图3-5所示，提醒与通知子系统覆盖创建/查询/删除提醒、调度触发与系统通知推送等路径；"
            "该用例强调事件驱动与状态一致性：提醒既要落库可追踪，也要在到期时可靠触达用户。"
        ),
        "3-6": (
            "如图3-6所示，截图与 OCR 子系统与桌宠等前端能力协同，覆盖采集、识别、检索与证据化引用；"
            "其关键价值在于把桌面行为轨迹结构化，为对话回答提供可溯源依据。"
        ),
        "3-7": (
            "如图3-7所示，主要加工链路将对话、提醒、截图与资料操作映射到统一的数据落库与外部依赖关系；"
            "该图突出“业务动作—持久化—外部服务”的对应关系，便于评审数据一致性与故障域。"
        ),
        "3-8": (
            "如图3-8所示，系统逻辑架构在桌面侧与后端侧分别展开关键模块，并标注主要通信方式（HTTP/WebSocket/IPC）；"
            "该视图用于从实现视角对齐第2章相关技术与第4章模块划分。"
        ),
        "3-9": (
            "如图3-9所示，分层架构进一步细化模块边界与依赖方向，强调渲染、主进程与后端服务的职责分离；"
            "这有助于在扩展功能时控制耦合面，避免跨层调用导致的维护成本上升。"
        ),
        "3-10": (
            "如图3-10所示，启动与 IPC 能力将通道注册、处理器挂载与前端 API 暴露串联为可复用的基础能力；"
            "该图解释为何即使配置加载异常，系统仍优先保证 IPC 可用，以避免界面层出现不可恢复错误。"
        ),
        "3-11": (
            "如图3-11所示，核心业务功能模块围绕智能对话、记忆、提醒、截图与资料夹形成闭环；"
            "模块之间通过主进程编排组合，体现“同一套底座支撑多业务场景”的设计取向。"
        ),
        "3-12": (
            "如图3-12所示，数据持久化与外部依赖模块明确 SQLite、调度、WebSocket 与 OCR 等组件的边界；"
            "该图为部署排错提供抓手：当某一外部依赖不可用时，可快速定位影响的业务链路。"
        ),
        "3-13": (
            "如图3-13所示，总架构 ER 图给出核心实体及其关联，体现对话、提醒、截图等业务对象的聚合关系；"
            "该模型与后端表结构相互对照，用于保证字段设计与查询路径一致。"
        ),
        "3-14": (
            "如图3-14所示，对话记忆相关表结构聚焦消息序列、会话标识与清理策略等字段；"
            "其设计目标是兼顾检索效率与隐私可控（如按消息粒度删除）。"
        ),
        "3-15": (
            "如图3-15所示，人设相关表结构支撑默认人设、覆盖人设与版本化文本；"
            "通过结构化存储可降低模型提示拼装阶段的歧义，并便于审计变更。"
        ),
        "3-16": (
            "如图3-16所示，提醒与截屏相关表结构覆盖计划任务、触发状态与 OCR 结果等关键字段；"
            "该结构支撑定时采集、状态查询与对话检索等多条链路。"
        ),
        "3-17": (
            "如图3-17所示，核心对话流程按时序展示从用户输入到工具调用与回复落库的交互；"
            "该时序强调多轮工具循环与异常降级路径，是验证系统可解释性与稳定性的重要依据。"
        ),
        "3-18": (
            "如图3-18所示，人设意图处理流程刻画自然语言更新人设时的抽取、校验与持久化步骤；"
            "其目的在于减少“口头承诺未落库”或“覆盖范围不清”等一致性风险。"
        ),
        "3-19": (
            "如图3-19所示，提醒意图处理流程突出时间解析、冲突校验与多端触达的衔接；"
            "该时序可用于对照第4章提醒模块实现与第5章相关测试结果。"
        ),
        "4-1": (
            "如图4-1所示，系统首页提供主要功能入口与导航结构，体现信息架构与交互路径；"
            "该界面作为用户首次接触系统的触点，其清晰度直接影响后续功能发现效率。"
        ),
        "4-2": (
            "如图4-2所示，基础对话界面展示消息流、输入区与状态反馈区域；"
            "界面布局遵循“连续阅读 + 明确操作”原则，以降低多轮对话的认知负担。"
        ),
        "4-3": (
            "如图4-3所示，对话历史页面支持浏览与定位历史消息，为记忆管理与排错提供可视化手段；"
            "该页面与本地持久化策略相对应，是验证记忆一致性的重要界面。"
        ),
        "4-4": (
            "如图4-4所示，提醒功能页面呈现提醒列表与关键字段（时间、标题、状态等）；"
            "其设计强调可扫描性与可操作性的平衡，便于用户快速管理日程类信息。"
        ),
        "4-5": (
            "如图4-5所示，截图轨迹页面用于展示采集记录与时间线，为 OCR 结果可追溯提供入口；"
            "该界面是“轨迹—证据—对话”链路中的关键观测点。"
        ),
        "4-6": (
            "如图4-6所示，截图区域框定交互用于减少无关像素对 OCR 的干扰；"
            "通过约束采集范围，可显著提升识别稳定性与后续检索命中率。"
        ),
        "4-7": (
            "如图4-7所示，截图内容识别结果以结构化方式呈现 OCR 文本与状态；"
            "该反馈不仅服务用户自查，也为模型侧工具调用提供可靠输入。"
        ),
        "4-8": (
            "如图4-8所示，AI 资料夹界面体现受控目录下的文件列表与操作入口；"
            "其边界与权限策略相对应，避免越权读写带来的安全风险。"
        ),
        "4-9": (
            "如图4-9所示，清空对话记录操作提供明确的危险操作路径与确认机制；"
            "该交互设计强调可逆性提示与状态同步，降低误删带来的数据损失。"
        ),
        "4-10": (
            "如图4-10所示，AI 自助保存资料的回复体现工具调用结果在对话中的呈现方式；"
            "其可用于核对“自然语言指令—文件系统动作—反馈文本”的一致性。"
        ),
    }


def _next_meaningful_sibling(el):
    x = el.getnext()
    while x is not None:
        if x.tag == qn("w:p"):
            t = _para_text(x).strip()
            if t:
                return x
        x = x.getnext()
    return None


def _already_has_explanation(after_cap_el) -> bool:
    n = _next_meaningful_sibling(after_cap_el)
    if n is None:
        return False
    t = _para_text(n)
    return t.startswith("如图") or t.startswith("表3-1") or t.startswith("综上")


def fix_table31_and_intro(doc: Document) -> None:
    body = doc.element.body
    tbl = None
    cap_el = None
    for el in body.iterchildren():
        if el.tag == qn("w:tbl"):
            nxt = el.getnext()
            if nxt is not None and nxt.tag == qn("w:p"):
                t = _para_text(nxt).strip().replace(" ", "")
                if t.startswith("表3-1"):
                    tbl, cap_el = el, nxt
                    break
    if tbl is None or cap_el is None:
        for el in body.iterchildren():
            if el.tag != qn("w:p"):
                continue
            t = _para_text(el).strip().replace(" ", "")
            if not t.startswith("表3-1"):
                continue
            nxt = el.getnext()
            if nxt is not None and nxt.tag == qn("w:tbl"):
                cap_el, tbl = el, nxt
                break
    if tbl is None or cap_el is None:
        return

    # Normalize caption
    from docx.text.paragraph import Paragraph

    cap_p = Paragraph(cap_el, doc)
    if cap_p.runs:
        cap_p.text = "表3-1 业务需求表"

    if cap_el.getnext() is tbl:
        pass
    else:
        cap_el.getparent().remove(cap_el)
        tbl.addprevious(cap_el)

    # Intro before caption (if not already present)
    prev = cap_el.getprevious()
    intro_txt = (
        "结合前述六类能力划分，将关键业务对象、约束条件与交付结果进行结构化汇总，并形成业务需求清单，如表3-1所示。"
    )
    need_intro = True
    if prev is not None and prev.tag == qn("w:p"):
        pt = _para_text(prev)
        if "如表3-1所示" in pt or "业务需求清单" in pt:
            need_intro = False
    if need_intro:
        h = Document()
        p = h.add_paragraph(intro_txt)
        p.paragraph_format.first_line_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.5
        for r in p.runs:
            _set_run_body_font(r)
        cap_el.addprevious(deepcopy(p._element))

    # Explanation after table
    explain = (
        "由表3-1可知，系统需求覆盖智能对话与记忆、提醒与通知、截图与 OCR、资料夹与导航等核心域；"
        "各需求条目与后文章节中的模块设计、接口划分及测试用例形成对应关系，可作为全稿一致性与可追溯性的基线。"
    )
    nxt = tbl.getnext()
    if nxt is not None and nxt.tag == qn("w:p") and "由表3-1可知" in _para_text(nxt):
        return
    _insert_after(body, tbl, [_body_paragraph(doc, explain)])


def normalize_figure_captions(doc: Document) -> None:
    from docx.text.paragraph import Paragraph

    for el in doc.element.body.iterchildren():
        if el.tag != qn("w:p"):
            continue
        raw = _para_text(el).strip()
        norm = _normalize_figure_caption_text(raw)
        if not norm:
            continue
        p = Paragraph(el, doc)
        if p.text.strip() != norm:
            p.text = norm


def insert_figure_explanations(doc: Document) -> None:
    body = doc.element.body
    expl = EXPLANATIONS()
    for el in list(body.iterchildren()):
        if el.tag != qn("w:p"):
            continue
        raw = _para_text(el).strip()
        m = FIG_CAP_RE.match(raw)
        if not m:
            continue
        key = f"{m.group(1)}-{m.group(2)}"
        text = expl.get(key)
        if not text:
            continue
        if _already_has_explanation(el):
            continue
        _insert_after(body, el, [_body_paragraph(doc, text)])


def insert_chapter_transitions(doc: Document) -> None:
    body = doc.element.body

    # Before 3.3: bridge from functional figures to NFR
    for el in body.iterchildren():
        if el.tag != qn("w:p"):
            continue
        t = _para_text(el).strip()
        if not (t.startswith("3.3") and "非功能" in t):
            continue
        prev = el.getprevious()
        bridge = (
            "综合上述业务流程刻画与用例边界，功能需求已在结构上与后续设计相衔接；"
            "以下进一步从性能、稳定性、安全性与可维护性等方面给出非功能需求分析。"
        )
        if prev is not None and prev.tag == qn("w:p") and "以下进一步从性能" in _para_text(prev):
            break
        _insert_before(body, el, [_body_paragraph(doc, bridge)])
        break

    # Before 3.5: bridge from architecture figures to data flow
    for el in body.iterchildren():
        if el.tag != qn("w:p"):
            continue
        t = _para_text(el).strip()
        if not (t.startswith("3.5") and "数据" in t):
            continue
        bridge = (
            "结合图3-7至图3-12的分层与依赖关系，系统在进程通信、模块边界与外部组件方面已形成一致的结构化视图；"
            "在此前提下，本节从数据流角度归纳核心业务链路，突出各链路的输入、处理、落库与反馈环节。"
        )
        prev = el.getprevious()
        if prev is not None and prev.tag == qn("w:p") and "在此前提下，本节从数据流角度" in _para_text(prev):
            break
        _insert_before(body, el, [_body_paragraph(doc, bridge)])
        break

    # Before 3.6: bridge from model/sequence figures to summary
    for el in body.iterchildren():
        if el.tag != qn("w:p"):
            continue
        t = _para_text(el).strip()
        if not (t.startswith("3.6") and "小结" in t):
            continue
        bridge = (
            "上述 ER 与时序图示从静态结构与动态交互两方面补充了需求分析的可验证材料；"
            "本节最后对第3章工作进行归纳，并为第4章详细设计与实现提供直接对照关系。"
        )
        prev = el.getprevious()
        if prev is not None and prev.tag == qn("w:p") and "本节最后对第3章工作进行归纳" in _para_text(prev):
            break
        _insert_before(body, el, [_body_paragraph(doc, bridge)])
        break

    # After UI intro "主要界面效果如下所示" add bridge if next is figure caption only
    for el in body.iterchildren():
        if el.tag != qn("w:p"):
            continue
        t = _para_text(el).strip()
        if t != "主要界面效果如下所示。":
            continue
        bridge = (
            "为便于对照第4章实现细节，本节给出代表性界面截图；"
            "下列图示侧重展示信息架构、关键操作路径以及与主进程能力的交互反馈。"
        )
        nxt = _next_meaningful_sibling(el)
        if nxt is not None and "下列图示侧重展示" in _para_text(nxt):
            break
        _insert_after(body, el, [_body_paragraph(doc, bridge)])
        break


def main() -> int:
    docs_dir = os.path.join(_paper_root(), "docs")
    path = _resolve_thesis_path(docs_dir)
    doc = Document(path)

    fix_table31_and_intro(doc)
    normalize_figure_captions(doc)
    insert_figure_explanations(doc)
    insert_chapter_transitions(doc)

    doc.save(path)
    print("Updated:", path)
    return 0


if __name__ == "__main__":
    sys.exit(main())
