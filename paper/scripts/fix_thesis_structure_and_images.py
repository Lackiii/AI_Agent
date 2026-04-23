# -*- coding: utf-8 -*-
"""
Re-structure chapter 3 figures, fix numbering, embed images from paper/imgs.
Source: thesis backup (no broken 3.5.1 merge). Writes the main 初稿 docx.

Run: python paper/scripts/fix_thesis_structure_and_images.py
Then (captions + explanations + transitions): python paper/scripts/enrich_thesis_figures_and_tables.py
"""
from __future__ import annotations

import os
import re
import shutil
import sys
from copy import deepcopy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
def _paper_root() -> str:
    return os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))


def _resolve_thesis_paths(docs_dir: str) -> tuple[str, str]:
    """Return (main_thesis_path, backup_path)."""
    main = backup = None
    for name in os.listdir(docs_dir):
        if not name.endswith(".docx") or name.startswith("~$"):
            continue
        if "2022433020102" not in name:
            continue
        low = name.lower()
        path = os.path.join(docs_dir, name)
        if "backup_before_api_merge" in low:
            backup = path
        else:
            main = path
    if not backup or not main:
        raise FileNotFoundError(f"Need thesis + backup in {docs_dir}, got main={main!r} backup={backup!r}")
    return main, backup


def _img(path: str) -> str:
    return os.path.join(_paper_root(), "imgs", path)


def _set_caption_style(run):
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _figure_elements(caption: str, image_file: str | None, width_inches: float = 5.6):
    """Return list of w:p OXML elements: optional image (centered) + caption (centered)."""
    h = Document()
    out = []
    if image_file and os.path.isfile(image_file):
        p_img = h.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.add_run().add_picture(image_file, width=Inches(width_inches))
        out.append(deepcopy(p_img._element))
    p_cap = h.add_paragraph(caption)
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p_cap.runs:
        _set_caption_style(p_cap.runs[0])
    out.append(deepcopy(p_cap._element))
    return out


def _body_children(doc: Document):
    return list(doc.element.body.iterchildren())


def _para_text(el) -> str:
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    ts = el.findall(".//w:t", namespaces=ns)
    return "".join(t.text or "" for t in ts)


def _insert_before(body, anchor_elm, elements: list):
    children = list(body.iterchildren())
    idx = children.index(anchor_elm)
    for i, el in enumerate(elements):
        body.insert(idx + i, el)


def _remove_elements(body, elements: list):
    for el in elements:
        body.remove(el)


def replace_figures_between_flow_intro_and_33(doc: Document):
    """Remove everything between '总体业务流程图' and '3.3 非功能需求分析', then insert 图3-1…3-6 from imgs."""
    body = doc.element.body
    children = _body_children(doc)
    intro_i = None
    i33 = None
    for i, el in enumerate(children):
        if el.tag != qn("w:p"):
            continue
        t = _para_text(el)
        if "总体业务流程图" in t and intro_i is None:
            intro_i = i
        if t.strip().startswith("3.3") and "非功能" in t:
            i33 = i
            break
    if intro_i is None or i33 is None:
        raise RuntimeError("Could not locate 总体业务流程图 block or 3.3 heading")
    intro_el = children[intro_i]
    anchor_33 = children[i33]
    to_remove = children[intro_i + 1 : i33]
    _remove_elements(body, to_remove)

    specs = [
        ("图3-1 应用启动（Electron 主链路）流程图", _img("应用启动（Electron 主链路）流程图.png")),
        ("图3-2 渲染层调用主进程流程图", _img("渲染层调用主进程流程图.png")),
        ("图3-3 模块分层关系图", _img("模块分层关系图.png")),
        ("图3-4 核心对话系统用例图", _img("核心对话系统用例图.png")),
        ("图3-5 提醒与通知子系统用例图", _img("提醒与通知子系统用例图.png")),
        ("图3-6 截图OCR与桌宠子系统用例图", _img("截图 OCR + 桌宠子系统用例图.png")),
    ]
    new_els: list = []
    for cap, fp in specs:
        new_els.extend(_figure_elements(cap, fp))
    _insert_before(body, anchor_33, new_els)


def insert_architecture_before_35(doc: Document):
    body = doc.element.body
    anchor = None
    for el in body.iterchildren():
        if el.tag != qn("w:p"):
            continue
        t = _para_text(el).strip()
        if t.startswith("3.5") and "数据" in t:
            anchor = el
            break
    if anchor is None:
        raise RuntimeError("Anchor 3.5 数据 not found")
    specs = [
        ("图3-7 主要加工与数据存储架构图", _img("主要加工与数据存储架构图.png")),
        ("图3-8 系统架构图", _img("系统架构图.png")),
        ("图3-9 整体分层架构模块图", _img("整体分层架构模块图.png")),
        ("图3-10 启动与IPC功能模块图", _img("启动 + IPC功能模块图.png")),
        ("图3-11 核心业务功能模块图", _img("核心业务功能模块图.png")),
        ("图3-12 数据持久化与外部依赖模块图", _img("数据持久化与外部依赖模块图.png")),
    ]
    intro = Document()
    p = intro.add_paragraph(
        "结合总体架构说明，给出加工链路、逻辑架构、分层模块、进程通信、业务功能与持久化依赖的结构化图示如下。"
    )
    block = [deepcopy(p._element)]
    for cap, fp in specs:
        block.extend(_figure_elements(cap, fp))
    _insert_before(body, anchor, block)


def insert_er_and_tables_before_36(doc: Document):
    body = doc.element.body
    anchor = None
    for el in body.iterchildren():
        if el.tag != qn("w:p"):
            continue
        t = _para_text(el).strip()
        if t.startswith("3.6") and "小结" in t:
            anchor = el
            break
    if anchor is None:
        raise RuntimeError("Anchor 3.6 not found")
    intro = Document()
    p = intro.add_paragraph("为支撑上述数据流，持久化结构以 SQLite 为核心，核心表结构示意如下。")
    block = [deepcopy(p._element)]
    specs = [
        ("图3-13 总架构ER图", _img("总架构ER图.png")),
        ("图3-14 对话记忆相关数据表结构", _img("对话记忆数据库表.png")),
        ("图3-15 人设相关数据表结构", _img("人设数据库表.png")),
        ("图3-16 提醒及截屏相关数据表结构", _img("提醒及截屏功能数据库表.png")),
    ]
    for cap, fp in specs:
        block.extend(_figure_elements(cap, fp))
    seq_intro = Document()
    p2 = seq_intro.add_paragraph("主要业务流程的时序关系如下所示。")
    block.append(deepcopy(p2._element))
    seqs = [
        ("图3-17 核心对话流程时序图", _img("核心对话流程时序图.png")),
        ("图3-18 人设意图处理时序图", _img("人设意图处理时序图.png")),
        ("图3-19 提醒意图处理时序图", _img("提醒意图处理时序图.png")),
    ]
    for cap, fp in seqs:
        block.extend(_figure_elements(cap, fp))
    _insert_before(body, anchor, block)


def insert_chapter4_ui_figures(doc: Document):
    """After first chunk of 4.1, insert representative UI screenshots."""
    body = doc.element.body
    anchor = None
    seen_41 = False
    for el in body.iterchildren():
        if el.tag != qn("w:p"):
            continue
        t = _para_text(el).strip()
        if re.match(r"^4\.1\b", t):
            seen_41 = True
            continue
        if seen_41 and "最小权限" in t and "设计原则" in t:
            anchor = el
            break
    if anchor is None:
        return
    intro = Document()
    p = intro.add_paragraph("主要界面效果如下所示。")
    block = [deepcopy(p._element)]
    specs = [
        ("图4-1 系统首页", _img("首页.png")),
        ("图4-2 基础对话界面", _img("基础对话.png")),
        ("图4-3 对话历史页面", _img("对话历史页面.png")),
        ("图4-4 提醒功能页面", _img("提醒功能页面.png")),
        ("图4-5 截图轨迹页面", _img("截图轨迹页面.png")),
        ("图4-6 截图区域框定", _img("截图范围框定.png")),
        ("图4-7 截图内容识别结果", _img("截图内容识别.png")),
        ("图4-8 AI资料夹界面", _img("AI资料夹样式.png")),
        ("图4-9 清空对话记录操作", _img("清空对话记录.png")),
        ("图4-10 AI自助保存资料回复", _img("AI自助保存资料回复.png")),
    ]
    for cap, fp in specs:
        block.extend(_figure_elements(cap, fp))
    nxt = anchor.getnext()
    if nxt is not None:
        _insert_before(body, nxt, block)
    else:
        for el in block:
            body.append(el)


def strip_duplicate_35_section_if_present(doc: Document):
    """Remove erroneous '3.5.1 …' heading and following empty rows / orphan table before 3.6."""
    body = doc.element.body
    children = _body_children(doc)
    start_i = None
    for i, el in enumerate(children):
        if el.tag != qn("w:p"):
            continue
        t = _para_text(el)
        if "3.5.1" in t and ("接口" in t or "模块交互" in t):
            start_i = i
            break
    if start_i is None:
        return
    end_i = None
    for j in range(start_i + 1, len(children)):
        el = children[j]
        if el.tag == qn("w:p"):
            t = _para_text(el).strip()
            if t.startswith("3.6"):
                end_i = j
                break
        if el.tag == qn("w:tbl") and end_i is None:
            continue
    if end_i is None:
        return
    to_remove = children[start_i:end_i]
    _remove_elements(body, to_remove)


def main() -> int:
    docs_dir = os.path.join(_paper_root(), "docs")
    main_path, backup_path = _resolve_thesis_paths(docs_dir)

    out_backup = main_path + ".before_structure_fix.docx"
    if not os.path.isfile(out_backup) and os.path.isfile(main_path):
        shutil.copy2(main_path, out_backup)

    shutil.copy2(backup_path, main_path)
    doc = Document(main_path)

    replace_figures_between_flow_intro_and_33(doc)
    insert_architecture_before_35(doc)
    insert_er_and_tables_before_36(doc)
    insert_chapter4_ui_figures(doc)
    strip_duplicate_35_section_if_present(doc)

    doc.save(main_path)
    print("Wrote:", main_path)
    print("Images from:", os.path.join(_paper_root(), "imgs"))
    return 0


if __name__ == "__main__":
    sys.exit(main())
